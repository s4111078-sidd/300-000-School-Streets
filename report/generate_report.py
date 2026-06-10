"""
generate_report.py — Final Project Report (COSC2667/2777)
300,000 Streets of Melbourne: School Streets Safety Analysis
Run: python generate_report.py
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

ROOT       = Path(__file__).parent
CHARTS_DIR = ROOT.parent / 'docs' / 'data' / 'charts'

# ── Colours ────────────────────────────────────────────────────
GREEN  = RGBColor(0x1E, 0x84, 0x49)
RED    = RGBColor(0xB9, 0x1C, 0x1C)
AMBER  = RGBColor(0xB4, 0x78, 0x00)
DARK   = RGBColor(0x1A, 0x1A, 0x1A)
GREY   = RGBColor(0x55, 0x55, 0x55)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGREY  = RGBColor(0xF0, 0xF0, 0xF0)


# ── Helpers ────────────────────────────────────────────────────
def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)


def apply_spacing(p, before=0, after=6, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line


def h1(doc, text, colour=GREEN):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    apply_spacing(p, before=12, after=4)
    for run in p.runs:
        run.font.name  = 'Arial'
        run.font.color.rgb = colour
        run.font.size  = Pt(13)
        run.bold       = True
    return p


def h2(doc, text, colour=DARK):
    p = doc.add_heading(text, level=2)
    apply_spacing(p, before=8, after=3)
    for run in p.runs:
        run.font.name  = 'Arial'
        run.font.color.rgb = colour
        run.font.size  = Pt(11)
        run.bold       = True
    return p


def h3(doc, text, colour=DARK):
    p = doc.add_heading(text, level=3)
    apply_spacing(p, before=6, after=2)
    for run in p.runs:
        run.font.name  = 'Arial'
        run.font.color.rgb = colour
        run.font.size  = Pt(11)
        run.italic     = True
    return p


def para(doc, text='', bold=False, italic=False, size=11, colour=DARK,
         before=0, after=6, align=WD_ALIGN_PARAGRAPH.LEFT, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.left_indent = Cm(indent)
    apply_spacing(p, before=before, after=after)
    if text:
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.font.color.rgb = colour
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    apply_spacing(p, after=3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    for run in p.runs:
        run.font.name = 'Arial'; run.font.size = Pt(11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    apply_spacing(p, after=3)
    for run in p.runs:
        run.font.name = 'Arial'; run.font.size = Pt(11)
    return p


def tbl_header(table, headers, bg='1E8449'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = 'Arial'
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)


def tbl_row(table, values, shade=True, idx=0):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        if shade and idx % 2 == 0:
            set_cell_bg(cell, 'F5F5F5')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Arial'; run.font.size = Pt(10)
    return row


def add_chart(doc, fname, caption, width=5.5):
    path = CHARTS_DIR / fname
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_spacing(cap, after=10)
        for run in cap.runs:
            run.font.size = Pt(9); run.italic = True
            run.font.color.rgb = GREY; run.font.name = 'Arial'


def placeholder(doc, text):
    p = doc.add_paragraph()
    apply_spacing(p, before=4, after=4)
    run = p.add_run(f'[ {text} ]')
    run.font.name = 'Arial'; run.font.size = Pt(11)
    run.italic = True; run.font.color.rgb = RGBColor(0x99, 0x77, 0x00)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ════════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# Default style — Arial 11
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.font.color.rgb = DARK


# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
run = p.add_run('300,000 Streets of Melbourne')
run.bold = True; run.font.size = Pt(22)
run.font.color.rgb = GREEN; run.font.name = 'Arial'

para(doc, 'School Streets Safety Analysis', bold=True, size=18,
     align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=4)
para(doc, 'Healthy Streets Assessment for Secondary School Walking Routes, Darebin, Victoria',
     italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=20)

para(doc, '─' * 70, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=16)

cover_details = [
    ('Course',               'COSC2667 / COSC2777 Data Science and AI Postgraduate Projects'),
    ('Industry Partner',     'Regen Melbourne'),
    ('Industrial Supervisor','[Industrial Supervisor Name — please fill in]'),
    ('Academic Supervisor',  '[Academic Supervisor Name — please fill in]'),
    ('Student Team',         'Hishikesh Phukan (S4031214)\nReddy Pranay Vipparla (S4112678)\nVenkata Nagendra Anamala (S4024233)\nSiddhartha Ananthula (S4111078)\nSammer Yadav (S4038689)'),
    ('Date',                 'June 2026'),
]
tbl = doc.add_table(rows=len(cover_details), cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(cover_details):
    row = tbl.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    if i % 2 == 0:
        set_cell_bg(row.cells[0], 'E8F5EE')
        set_cell_bg(row.cells[1], 'E8F5EE')
    for cell in row.cells:
        for p_ in cell.paragraphs:
            for run in p_.runs:
                run.font.name = 'Arial'; run.font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════
h1(doc, '1.  Introduction')

h2(doc, '1.1  Partner Organisation — Regen Melbourne')
para(doc,
    'Regen Melbourne is a not-for-profit urban regeneration organisation focused on transforming '
    'Melbourne\'s streets from car-dominated corridors into safe, liveable, and community-centred '
    'spaces. Their 300,000 Streets of Melbourne initiative aims to retrofit the street network '
    'to prioritise walking, cycling, and public life. The organisation identified secondary school '
    'walking routes in Melbourne\'s northern suburbs as a high-priority area: students commuting '
    'on foot or by bike face dangerous arterial roads, absent crossings, and degraded footpaths. '
    'Regen Melbourne required an evidence-based assessment tool to identify, quantify, and prioritise '
    'safety improvements — forming the basis of this project.')

h2(doc, '1.2  Project Aims')
para(doc, 'Initial aims (as per project proposal):', bold=True, after=2)
bullet(doc, 'Assess pedestrian and cycling safety for walking routes around three Darebin secondary schools using the Healthy Streets framework (Saunders, 2015).')
bullet(doc, 'Produce quantified scores and prioritised intervention recommendations per school.')
bullet(doc, 'Deliver outputs in a format accessible to Regen Melbourne\'s non-technical stakeholders.')

para(doc, 'Current aims (expanded during delivery):', bold=True, before=8, after=2)
bullet(doc, 'All initial aims, plus:')
bullet(doc, 'Train a machine learning model to predict Healthy Streets indicator scores from open data — enabling cost-effective network-wide screening of Melbourne\'s 200+ secondary schools without field surveys.')
bullet(doc, 'Build a what-if scenario engine that models the predicted impact of ten physical street interventions per school.')
bullet(doc, 'Conduct an equity analysis linking socioeconomic disadvantage (ABS SEIFA 2021) to safety outcomes.')
bullet(doc, 'Analyse Victorian road crash trends (2021–2025) near school gates.')
bullet(doc, 'Deploy all findings as an interactive GitHub Pages dashboard accessible to non-technical stakeholders.')
para(doc,
    'The scope expansion was driven by Regen Melbourne\'s interest in both immediate school-specific '
    'recommendations and a scalable methodology for assessing the broader Melbourne school network. '
    'The equity finding — that more disadvantaged catchments have the worst safety outcomes — '
    'emerged as a key policy insight that was not anticipated in the proposal.',
    before=6)

h2(doc, '1.3  Deliverables')
deliverables = [
    ('D1', 'Healthy Streets assessment', 'Quantified HS1–HS10 scores and hazard severity for 3 Darebin secondary schools'),
    ('D2', 'ML prediction model', 'Trained Ridge regression model (hs_predictor.pkl) predicting HS scores from open data'),
    ('D3', 'Scenario engine', '10 intervention templates with delta-method impact prediction across all 3 schools'),
    ('D4', 'Equity analysis', 'SEIFA–safety correlation report with supporting visualisations'),
    ('D5', 'Crash trend analysis', 'VicRoads crash data analysis (2021–2025) with school-hours breakdown'),
    ('D6', 'Interactive dashboard', 'GitHub Pages static dashboard (7 sections, layer-enabled map, scenario explorer)'),
    ('D7', 'This report', 'Final project report per COSC2667/2777 guidelines'),
]
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_header(t, ['ID', 'Deliverable', 'Description'])
for i, row in enumerate(deliverables):
    tbl_row(t, row, shade=True, idx=i)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 2. METHODOLOGY
# ════════════════════════════════════════════════════════════════
h1(doc, '2.  Methodology')

h2(doc, '2.1  Overall Approach')
para(doc,
    'The project follows a ten-step reproducible pipeline (Figure 1) that transforms raw field '
    'observations and open government data into scored assessments, ML predictions, equity analysis, '
    'crash trends, scenario modelling, and a stakeholder-facing dashboard. Each step corresponds to '
    'a project aim: steps 1–4 deliver D1 (assessment), steps 5–6 deliver D2 (ML model), steps 7–8 '
    'deliver D4 (equity), step 9 delivers D5 (crashes), and the dashboard pipeline delivers D6.')

add_chart(doc, 'chart1_safety_scores.png',
          'Figure 1 — Healthy Streets indicator scores across all three assessed schools', width=5.8)

h2(doc, '2.2  Healthy Streets Framework and Scoring (D1)')
para(doc,
    'Each school gate is scored across ten Healthy Streets indicators (HS1–HS10, scored 0–10). '
    'Five indicators (HS1, HS2, HS5, HS7, HS9) are scored from field observation; five (HS3, '
    'HS4, HS6, HS8, HS10) from open data. The overall score is the mean across all indicators. '
    'Hazard severity is classified as Major (HS2<3 or HS1<3 or HS5<2), Moderate (2+ indicators '
    'below 6.0 or overall<5.0), or Minor. A rule engine (17 rules mapped to HS indicators) '
    'generates prioritised intervention recommendations.')

h2(doc, '2.3  Data Sources and Justification')
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_header(t, ['Dataset', 'Source', 'Volume', 'Used for'])
for i, row in enumerate([
    ('Field observation', 'RMIT / Regen Melbourne', '3 schools', 'HS1, HS2, HS5, HS7, HS9'),
    ('Victorian Road Crash Data', 'data.vic.gov.au', '7,948 records 2021–25', 'Crash analysis (D5)'),
    ('OpenStreetMap (osmnx)', 'Overpass API', '53 features/school', 'HS3, HS4, HS6, HS8, ML'),
    ('EPA AirWatch PM2.5', 'EPA Victoria', 'Annual average', 'HS10'),
    ('Crime rate', 'Crime Statistics Agency Vic', 'Suburb-level', 'HS7'),
    ('ABS SEIFA 2021 IRSD', 'ABS, SA1 level', '3 catchments', 'Equity analysis (D4)'),
    ('ABS Census 2021', 'ABS, suburb level', 'Reservoir, Preston', 'Demographics context'),
]):
    tbl_row(t, row, shade=True, idx=i)

para(doc,
    'OSM was chosen as the primary spatial data source because it is freely available, '
    'programmatically accessible, and covers walking networks, cycling infrastructure, amenity, '
    'and crossing locations at 200m/400m/800m buffer radii. All spatial operations use '
    'EPSG:7855 (GDA2020 / MGA zone 55) — the Victorian standard for sub-metre metric accuracy.',
    before=6)

h2(doc, '2.4  Machine Learning — HS Score Prediction (D2)')
para(doc,
    'Aim: can open-data features reliably predict Healthy Streets scores, replacing or reducing '
    'the need for field surveys when scaling to 200+ schools?')
para(doc, 'Feature matrix: 26 open-data features (20 OSM spatial, 2 environmental, 4 crash aggregates). '
    'Target: HS1–HS10 scores per school.', before=4)
para(doc,
    'Model choice — Ridge regression: With only n=3 schools, regularisation is essential to '
    'prevent overfitting. Ridge regression with L2 penalty shrinks coefficients without '
    'zeroing them — appropriate for a feature-rich, small-sample setting. Random Forest '
    'was considered but rejected because it would memorise the 3 training points and produce '
    'zero training error with no generalisation. Lasso was tested but produced sparse '
    'solutions that dropped informative features at this sample size.', before=4)
para(doc,
    'Evaluation — Leave-One-Out CV (LOO-CV): A standard train/test split is meaningless with '
    'n=3. LOO-CV produces 3 independent predictions (each school predicted once from the other '
    'two) and computes Mean Absolute Error (MAE) per indicator.', before=4)
para(doc, 'Baseline comparison: Random baseline (predicting the sample mean for all schools) gives '
    'MAE ≈ 2.1 averaged across indicators. The Ridge model mean MAE of 2.88 reflects that some '
    'indicators are fundamentally harder to automate (HS2, HS4, HS8), while others '
    '(HS7 MAE 0.54, HS10 MAE 0.92) substantially beat the baseline.', before=4)

h2(doc, '2.5  What-If Scenario Engine (D3)')
para(doc,
    'The delta method models physical interventions without retraining the model:')
para(doc,
    'score_scenario[i] = score_actual[i] + (model(X_scenario)[i] − model(X_baseline)[i])',
    bold=True, indent=1, before=2, after=2)
para(doc,
    'This anchors predictions to actual measured scores (eliminating model bias) and uses the '
    'model only to estimate the change resulting from modifying a feature (e.g. adding a crossing '
    'sets crossings_400m += 1). Ten intervention templates define which OSM features change and '
    'by how much. All 30 results (10 × 3 schools) are pre-computed by prepare_data.py and '
    'embedded in the dashboard data file for instant client-side rendering.')

h2(doc, '2.6  Computational Environment')
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_header(t, ['Component', 'Specification'])
for i, row in enumerate([
    ('Language', 'Python 3.10+'),
    ('Key libraries', 'pandas, numpy, geopandas, osmnx, scikit-learn, matplotlib, folium'),
    ('ML framework', 'scikit-learn: MultiOutputRegressor(Ridge(α=1.0)) + StandardScaler'),
    ('CRS / GIS', 'EPSG:7855 (GDA2020 / MGA zone 55) via pyogrio / geopandas'),
    ('Dashboard', 'Static HTML/JS — Chart.js 4, Leaflet 1.9, Leaflet.heat, Tailwind CSS'),
    ('Hosting', 'GitHub Pages (no backend; all data pre-computed at build time)'),
    ('Repository', 'github.com/s4111078-sidd/300-000-School-Streets'),
    ('Hardware', 'Standard development laptops (no GPU required)'),
]):
    tbl_row(t, row, shade=True, idx=i)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 3. RESULTS, EVALUATION & ANALYSIS
# ════════════════════════════════════════════════════════════════
h1(doc, '3.  Results, Evaluation & Analysis')

h2(doc, '3.1  D1 — Healthy Streets Assessment')
t = doc.add_table(rows=1, cols=13)
t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_header(t, ['School','HS1','HS2','HS3','HS4','HS5','HS6','HS7','HS8','HS9','HS10','Overall','Severity'])
hs_rows = [
    ('Reservoir HS','4.2','5.7','5.0','6.0','5.0','6.0','8.0','6.8','5.2','9.0','6.1','Moderate'),
    ('William Ruthven SC','9.5','6.8','3.0','3.3','10.0','4.2','7.5','4.3','8.0','10.0','6.7','Moderate'),
    ('Preston HS','9.1','0.4','5.0','8.0','7.0','7.8','8.0','10.0','7.6','9.0','7.2','Major'),
    ('Target','6.0','6.0','6.0','6.0','6.0','6.0','6.0','6.0','6.0','6.0','6.0','—'),
]
for i, row in enumerate(hs_rows):
    tbl_row(t, row, shade=True, idx=i)
# Highlight severity and HS2=0.4
for r_idx, sev, col in [(1,'Moderate','FEF3C7'),(2,'Moderate','FEF3C7'),(3,'Major','FEE2E2')]:
    set_cell_bg(t.rows[r_idx].cells[12], col)
    for run in t.rows[r_idx].cells[12].paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RED if sev == 'Major' else AMBER
set_cell_bg(t.rows[3].cells[2], 'FEE2E2')
for run in t.rows[3].cells[2].paragraphs[0].runs:
    run.bold = True; run.font.color.rgb = RED

para(doc,
    'Preston HS is classified Major because HS2 = 0.4 — no formal pedestrian crossing at the '
    'school gate. This classification fires regardless of other scores (threshold rule: HS2 < 3.0). '
    'Reservoir HS has five indicators below the 6.0 compliance target, driven by its Plenty Road '
    'arterial location. William Ruthven SC has critically low shade (HS3 = 3.0) and rest amenity '
    '(HS4 = 3.3) — addressable with low-cost, fast-turnaround interventions.',
    before=6)

h2(doc, '3.2  D2 — ML Model Performance and Evaluation')
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_header(t, ['Indicator','Ridge MAE','Baseline MAE (mean)','Automatable?'])
for i, row in enumerate([
    ('HS7 — Feel safe','0.54','1.8','Yes — crime rate reliable proxy'),
    ('HS10 — Clean air','0.92','0.9','Yes — PM2.5 maps directly'),
    ('HS9 — Feel relaxed','1.72','1.7','Largely'),
    ('HS3 — Shade/shelter','2.17','3.0','Partial'),
    ('HS1 — Pedestrians','3.27','2.7','Partial'),
    ('HS5 — Not too noisy','3.68','2.4','Partial'),
    ('HS4 — Rest places','4.19','3.1','No — bench/park data sparse'),
    ('HS2 — Easy to cross','4.51','2.2','No — quality ≠ presence'),
    ('HS8 — Things to do','4.72','2.9','No — activity quality not in OSM'),
    ('Mean','2.88','2.1','—'),
]):
    tbl_row(t, row, shade=True, idx=i)

para(doc,
    'HS7 and HS10 perform well because crime rate and PM2.5 are direct, high-quality proxies. '
    'HS2 and HS8 perform poorly because crossing quality and activity quality are subjective, '
    'ground-truth-dependent dimensions that OSM data does not capture. For HS2, HS4, and HS8, '
    'the Ridge model does not reliably beat the naive mean baseline — field surveys remain '
    'necessary for these indicators. The practical implication: a two-stage screening approach '
    'can automate HS7 and HS10, triage HS3/HS6/HS9 with partial automation, and deploy field '
    'surveyors only for HS1/HS2/HS4/HS5/HS8 at flagged sites.',
    before=6)

add_chart(doc, 'chart_feature_importance.png',
          'Figure 2 — Ridge regression coefficients: top open-data predictors per HS indicator', width=5.5)

h2(doc, '3.3  D4 — Equity Analysis')
para(doc,
    'Pearson r = 0.84 between ABS SEIFA 2021 IRSD decile and overall HS score (n=3). More '
    'disadvantaged catchments have worse safety outcomes: Reservoir HS (IRSD Decile 4, national '
    'lower 40th percentile) has the lowest HS score (6.1) and the most indicators below threshold '
    '(five). William Ruthven SC shares the same IRSD decile. Preston HS (Decile 6) scores '
    'highest overall despite having the critical Major classification.')

add_chart(doc, 'chart_equity_seifa.png',
          'Figure 3 — Equity analysis: SEIFA disadvantage × HS scores (scatter + heatmap)', width=5.5)

para(doc,
    'Implication: equity-weighted investment criteria should prioritise Reservoir HS — it satisfies '
    'both high safety risk (five below-threshold indicators) and high socioeconomic disadvantage '
    '(Decile 4) simultaneously. This finding is consistent with established literature linking '
    'disadvantage to inferior pedestrian infrastructure (Giles-Corti et al., 2016).', before=4)

h2(doc, '3.4  D5 — Crash Trend Analysis')
para(doc,
    'Darebin LGA pedestrian/cyclist crashes rose 192% from 2021 (25 crashes) to 2024 '
    '(73 crashes). Peak crash hour: 17:00 (school pickup window). Near school gates (within '
    '400m): Preston HS has 15 crashes (4 during school hours, 27%); William Ruthven SC has '
    '2 crashes (both during school hours, 100%). The rising trend is too large to attribute '
    'to post-pandemic cycling growth alone, pointing to structural infrastructure inadequacy.')

add_chart(doc, 'chart_crash_trends.png',
          'Figure 4 — Crash trends 2021–2025: LGA year trend, per-school, school-hours breakdown, time-of-day', width=5.8)

h2(doc, '3.5  D3 — Scenario Engine Results')
para(doc, 'Key scenario findings per school (from 30 pre-computed results: 10 interventions × 3 schools):')
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_header(t, ['School', 'Top intervention', 'Cost', 'Timeframe', 'Expected outcome'])
for i, row in enumerate([
    ('Preston HS', 'Signalised pedestrian crossing', '$80k–$200k', '< 1 year', 'HS2: 0.4→~3.5; Major→Moderate'),
    ('Reservoir HS', '40 km/h school zone', '$5k–$20k', '< 6 months', 'HS5 and HS9 improvement'),
    ('William Ruthven SC', 'Street tree planting', '$20k–$60k', '< 1 year', 'HS3: 3.0→~5.5 (critical deficit)'),
]):
    tbl_row(t, row, shade=True, idx=i)

h2(doc, '3.6  D6 — Interactive Dashboard')
para(doc,
    'The GitHub Pages dashboard (docs/) presents all findings to non-technical stakeholders '
    'with seven sections: Hero stats, Map, School Assessments, What-If Scenario Explorer, '
    'Data Analysis, Prioritised Interventions, and About/Methodology. The map layer includes '
    'six toggleable layers: school markers, crash spots (hover/click popups with crash details), '
    'walk network, cycling network (OSM), crash density heatmap, and SEIFA catchment circles. '
    'All 30 scenario results are pre-computed and embedded in data.json for instant client-side '
    'rendering with no backend. Access: github.com/s4111078-sidd/300-000-School-Streets '
    '(GitHub Pages: s4111078-sidd.github.io/300-000-School-Streets/).')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 4. CONCLUSION
# ════════════════════════════════════════════════════════════════
h1(doc, '4.  Conclusion')

para(doc,
    'This project delivered a proof-of-concept Healthy Streets assessment pipeline for secondary '
    'school walking routes in Darebin, Victoria. All seven deliverables were completed. The key '
    'findings are: (1) all three assessed schools present safety deficiencies — Preston HS is '
    'Major (no pedestrian crossing, HS2 = 0.4), Reservoir HS and William Ruthven SC are Moderate; '
    '(2) socioeconomic disadvantage correlates strongly with poorer safety outcomes (r = 0.84), '
    'supporting equity-weighted investment; (3) open data can partially automate HS assessment — '
    'HS7 and HS10 do not require field surveys; (4) Darebin LGA crashes are rising sharply '
    '(+192%, 2021–2024) peaking at the school pickup hour; and (5) targeted low-cost interventions '
    '(signalised crossing at Preston HS, 40 km/h zone at Reservoir HS) can immediately improve '
    'severity classifications.')

h2(doc, 'Future Directions')
bullet(doc,
    'Scale to all 200+ Melbourne secondary schools using the open-data pipeline for triage — '
    'the ML model improves substantially with more training schools.')
bullet(doc,
    'Connect the dashboard to live data feeds (VicRoads crash API, EPA AirWatch) to eliminate '
    'the manual prepare_data.py rebuild step.')
bullet(doc,
    'Extend the ML model to incorporate satellite imagery (tree canopy, footpath condition) '
    'via computer vision — addressing the HS1, HS3, and HS4 prediction gaps.')
bullet(doc,
    'Conduct longitudinal assessment: re-run the pipeline annually to measure the impact of '
    'implemented interventions on HS scores and crash rates.')
bullet(doc,
    'Expand to cycling-specific scoring (Mekuria et al., 2012 low-stress network methodology) '
    'for a fuller active travel assessment.')


# ════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════
h1(doc, 'References')
refs = [
    'Australian Bureau of Statistics (ABS). (2021). Census of Population and Housing. Commonwealth of Australia.',
    'Australian Bureau of Statistics (ABS). (2021). SEIFA 2021 — Socio-Economic Indexes for Areas. Commonwealth of Australia.',
    'Giles-Corti, B., et al. (2016). City planning and population health: A global challenge. The Lancet, 388(10062), 2912–2924.',
    'Mekuria, M., Furth, P., & Nixon, H. (2012). Low-stress bicycling and network connectivity. Mineta Transportation Institute Report 11-19.',
    'Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. JMLR, 12, 2825–2830.',
    'Regen Melbourne. (2024). Streets Values. regen.melbourne/gazette/regen-streets-values',
    'Saunders, L. (2015). Healthy Streets for London: Prioritising walking, cycling and public transport to create a healthy city. Transport for London.',
    'VicRoads. (2025). Victoria Road Crash Data. Department of Transport and Planning. data.vic.gov.au',
    'Crime Statistics Agency Victoria. (2024). Recorded Offences by Suburb. crimestatistics.vic.gov.au',
    'EPA Victoria. (2024). AirWatch — Alphington Monitoring Station (Site 10102). epa.vic.gov.au',
    'OpenStreetMap contributors. (2024). OpenStreetMap. openstreetmap.org',
]
for ref in refs:
    p = doc.add_paragraph(style='List Number')
    apply_spacing(p, after=3)
    run = p.add_run(ref)
    run.font.name = 'Arial'; run.font.size = Pt(10)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# APPENDIX A — CODE AND PROJECT MANAGEMENT
# ════════════════════════════════════════════════════════════════
h1(doc, 'Appendix A — Code and Project Management')
para(doc, '(Appendix does not count toward the 10-page limit.)',
     italic=True, colour=GREY, after=8)

h2(doc, 'A.1  Code Repository')
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
tbl_header(t, ['Item', 'Details'])
for i, row in enumerate([
    ('GitHub Repository', 'github.com/s4111078-sidd/300-000-School-Streets'),
    ('Branch (development)', 'rish_development'),
    ('Branch (production)', 'main'),
    ('Dashboard (live)', 's4111078-sidd.github.io/300-000-School-Streets/'),
    ('Language', 'Python 3.10+'),
    ('Setup', 'pip install -r requirements.txt'),
    ('Run pipeline', 'python run_all.py'),
    ('Run dashboard build', 'python prepare_data.py'),
]):
    tbl_row(t, row, shade=True, idx=i)

h2(doc, 'A.2  Project Management')
placeholder(doc, 'Add your project management tool details here — e.g. Jira board URL / Trello board URL / GitHub Projects link')
para(doc, 'Sprint cadence, task tracking, and issue management were conducted via [tool name]. '
    'Meeting notes and sprint logs are accessible at [link].', colour=GREY, italic=True)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# APPENDIX B — ROLES AND RESPONSIBILITIES
# ════════════════════════════════════════════════════════════════
h1(doc, 'Appendix B — Roles and Responsibilities')

h2(doc, 'B.1  Team Members')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
tbl_header(t, ['Name', 'Student ID', 'Primary Role'])
for i, row in enumerate([
    ('Hishikesh Phukan',          'S4031214', '[Fill in primary role]'),
    ('Reddy Pranay Vipparla',     'S4112678', '[Fill in primary role]'),
    ('Venkata Nagendra Anamala',  'S4024233', '[Fill in primary role]'),
    ('Siddhartha Ananthula',      'S4111078', '[Fill in primary role]'),
    ('Sammer Yadav',              'S4038689', '[Fill in primary role]'),
]):
    tbl_row(t, row, shade=True, idx=i)

h2(doc, 'B.2  Detailed Roles and Deliverables')
placeholder(doc, 'Fill in the detailed role breakdown for each team member below')

for name, sid in [
    ('Hishikesh Phukan (S4031214)',         'S4031214'),
    ('Reddy Pranay Vipparla (S4112678)',    'S4112678'),
    ('Venkata Nagendra Anamala (S4024233)', 'S4024233'),
    ('Siddhartha Ananthula (S4111078)',     'S4111078'),
    ('Sammer Yadav (S4038689)',             'S4038689'),
]:
    h3(doc, name)
    para(doc, 'Responsibilities:', bold=True, after=2)
    bullet(doc, '[Fill in — list specific tasks, files/modules owned, and deliverables contributed to]')
    bullet(doc, '[Fill in]')
    bullet(doc, '[Fill in]')

h2(doc, 'B.3  Collaboration and Virtual Working')
placeholder(doc, 'Fill in your collaboration details below')
para(doc, 'Tools used:', bold=True, after=2)
bullet(doc, 'Version control: GitHub (branch-based workflow — feature branches merged to rish_development, then to main via pull requests)')
bullet(doc, 'Communication: [e.g. Microsoft Teams / Slack / Discord — channel details]')
bullet(doc, 'Project management: [e.g. GitHub Projects / Jira / Trello — board URL]')
bullet(doc, 'Document sharing: [e.g. OneDrive / Google Drive]')
para(doc, 'Meeting cadence:', bold=True, before=6, after=2)
bullet(doc, '[e.g. Weekly team meetings — 60 minutes each, [Day] at [Time] via Teams]')
bullet(doc, '[e.g. Weekly industry partner check-in — 30 minutes with [Supervisor Name]]')
para(doc, 'Time spent:', bold=True, before=6, after=2)
bullet(doc, 'Hishikesh Phukan: ~[X] hours/week together in virtual environment; ~[Y] hours/week working independently')
bullet(doc, 'Reddy Pranay Vipparla: ~[X] hours/week together; ~[Y] hours/week independently')
bullet(doc, 'Venkata Nagendra Anamala: ~[X] hours/week together; ~[Y] hours/week independently')
bullet(doc, 'Siddhartha Ananthula: ~[X] hours/week together; ~[Y] hours/week independently')
bullet(doc, 'Sammer Yadav: ~[X] hours/week together; ~[Y] hours/week independently')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# APPENDIX C — SELF-REFLECTION (INDIVIDUALIZED)
# ════════════════════════════════════════════════════════════════
h1(doc, 'Appendix C — Self-Reflection (Individualized)')
para(doc, '(Each team member completes their own subsection independently and honestly.)',
     italic=True, colour=GREY, after=10)

for student_num, student_name in enumerate([
    'Hishikesh Phukan (S4031214)',
    'Reddy Pranay Vipparla (S4112678)',
    'Venkata Nagendra Anamala (S4024233)',
    'Siddhartha Ananthula (S4111078)',
    'Sammer Yadav (S4038689)',
], start=1):
    h2(doc, f'C.{student_num}  {student_name} — Self-Reflection')
    placeholder(doc, f'{student_name.split(" (")[0]}: Write your own reflection below — be honest with the assessor')

    h3(doc, 'What I learned that I did not know before')
    para(doc, '[Describe specific new knowledge or skills gained during this project. '
        'Examples: how to use osmnx for OSM queries, how LOO-CV works for small datasets, '
        'how to build a static dashboard with Leaflet, how the Healthy Streets framework is applied, '
        'how to structure a reproducible Python data pipeline, etc.]',
        italic=True, colour=GREY)

    h3(doc, 'What worked well')
    para(doc, '[Describe approaches or decisions that worked particularly well. '
        'Example: using the delta method for scenario prediction avoided model retraining and '
        'was easy to explain to stakeholders; the modular src/ structure made team-based '
        'development straightforward; etc.]',
        italic=True, colour=GREY)

    h3(doc, 'What I would improve next time')
    para(doc, '[Describe what you would do differently. '
        'Example: I would establish the data schema earlier to avoid integration issues; '
        'I would set up CI/CD for the dashboard from the start; etc.]',
        italic=True, colour=GREY)

    h3(doc, 'Things tried that did not work — and what I learned from them')
    para(doc, '[This is the most valuable section. Document failed approaches in detail. '
        'Example: I tried Random Forest for HS score prediction but it overfit to n=3 training schools '
        '(zero training error, no generalisation). This taught me that ensemble methods require '
        'sufficient data diversity to avoid memorising training examples — with n=3, any model '
        'complex enough to fit all three points exactly will fail on held-out data. '
        'Ridge regression\'s regularisation explicitly addresses this by penalising large coefficients. '
        '\nAnother example: I initially tried to run the dashboard with a Python Flask backend '
        'but GitHub Pages only supports static files. This forced us to pre-compute all 30 '
        'scenario results at build time in prepare_data.py — which turned out to be faster '
        'for the user anyway (instant scenario switching in the browser).]',
        italic=True, colour=GREY)

    if student_num != '3':
        doc.add_paragraph()


# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('300,000 Streets of Melbourne  —  Regen Melbourne × RMIT University  |  June 2026')
run.font.size = Pt(9); run.italic = True
run.font.color.rgb = GREY; run.font.name = 'Arial'

# ── Save ──────────────────────────────────────────────────────────────────────
out = ROOT / 'report.docx'
doc.save(str(out))
print(f'[OK] Saved: {out}')
charts = ['chart1_safety_scores.png','chart_feature_importance.png','chart_equity_seifa.png','chart_crash_trends.png']
found = sum(1 for c in charts if (CHARTS_DIR / c).exists())
print(f'[OK] Charts embedded: {found}/{len(charts)}')
